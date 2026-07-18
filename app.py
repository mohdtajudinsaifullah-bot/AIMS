import streamlit as st
import os
import time
import re
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_pinecone import PineconeVectorStore
from langchain_community.tools import DuckDuckGoSearchRun
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from io import BytesIO

# 1. Buka kunci (.env)
load_dotenv(override=True)
try:
    api_key_openai = st.secrets["OPENAI_API_KEY"] # Untuk Streamlit Cloud
    api_key_pinecone = st.secrets["PINECONE_API_KEY"]
except:
    api_key_openai = os.getenv("OPENAI_API_KEY")  # Untuk Laptop Lokal
    api_key_pinecone = os.getenv("PINECONE_API_KEY")

os.environ["OPENAI_API_KEY"] = api_key_openai
os.environ["PINECONE_API_KEY"] = api_key_pinecone

# --- FUNGSI AUTO-RETRY (VERSI OPENAI) ---
class ResponsAI:
    def __init__(self, text):
        self.text = text

def cuba_jana_ai(prompt_teks):
    senarai_model = ["gpt-4o-mini", "gpt-3.5-turbo", "gpt-4o"] 
    max_cuba = 3 
    for cubaan in range(max_cuba):
        for model_ai in senarai_model:
            try:
                llm = ChatOpenAI(model=model_ai, temperature=0.3)
                respons = llm.invoke(prompt_teks)
                return ResponsAI(respons.content) # Bungkus supaya serasi dengan '.text'
            except Exception as e:
                ralat = str(e)
                if "429" in ralat or "RateLimit" in ralat:
                    time.sleep(5) 
                    continue 
                elif "404" in ralat or "not found" in ralat.lower():
                    continue 
                else:
                    raise e 
    raise Exception("Pelayan OpenAI sesak atau baki kuota tidak mencukupi. Sila cuba lagi.")

# --- FUNGSI BINA FAIL WORD AP & PENGURUSAN ---
def bina_fail_word(teks_ai, tajuk_dokumen, metadata, is_pengurusan=False, is_arahan_amalan=False):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12) 
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(12)

    # --- 1. PEMBINAAN KEPALA SURAT ---
    if not is_pengurusan:
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(0)
        run_p1 = p1.add_run(metadata['pihak1'].upper())
        run_p1.bold = True

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        run_lwn = p2.add_run("lwn")
        run_lwn.italic = True

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(24)
        run_p3 = p3.add_run(metadata['pihak2'].upper())
        run_p3.bold = True

        p4 = doc.add_paragraph(f"[{metadata['mahkamah']} {metadata['negeri']}, {metadata['hakim']}, pada {metadata['tarikh']}]")
        p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p4.paragraph_format.space_after = Pt(0)
        
        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p5.paragraph_format.space_after = Pt(12)
        p5.add_run(f"[Kes No. {metadata['nokes']}]\n[{metadata['jenis_p']}]")

        if metadata['peguam']:
            p6 = doc.add_paragraph(f"Peguam Syarie: {metadata['peguam']}")
            p6.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p6.paragraph_format.space_after = Pt(24)
    else:
        p_tajuk = doc.add_paragraph()
        p_tajuk.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tajuk.paragraph_format.space_after = Pt(24)
        run_t = p_tajuk.add_run(tajuk_dokumen.upper())
        run_t.bold = True

    # --- 2. PEMBERSIHAN SAMPAH MARKDOWN AI & HALUSINASI ---
    teks_bersih = teks_ai.replace('**', '').replace('##', '').replace('#', '').replace('*', '').replace('>', '')
    lines = teks_bersih.split('\n')
    
    senarai_hitam_header = ["DI DALAM MAHKAMAH", "DALAM NEGERI", "ANTARA", "DENGAN", "DI HADAPAN", "PERMOHONAN NO", "ALASAN PENGHAKIMAN"]
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if any(line.upper().startswith(haram) for haram in senarai_hitam_header) and not is_pengurusan:
            i += 1
            continue

        if line.upper() == "KEPUTUSAN" or line.upper() == "5. KEPUTUSAN" or line.upper() == "4. KEPUTUSAN":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            run = p.add_run("KEPUTUSAN")
            run.bold = True
            
        elif line.upper().startswith("SETELAH") and not is_pengurusan:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            words = line.split(" ", 1)
            run_setelah = p.add_run(words[0].upper() + " ")
            run_setelah.bold = True
            run_setelah.underline = True
            if len(words) > 1:
                p.add_run(words[1])

        elif line.upper() in ["PERMOHONAN", "FAKTA KES", "ULASAN MAHKAMAH", "UNDANG-UNDANG YANG DIPAKAI"] or re.match(r'^\d+\.\s+[A-Z\s]+$', line.upper()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            run = p.add_run(line.upper())
            run.bold = True
            
        elif line.startswith('|') and not is_arahan_amalan:
            table_data = []
            while i < len(lines) and line.startswith('|'):
                cells = [c.strip() for c in lines[i].split('|') if c.strip()]
                if cells: table_data.append(cells)
                i += 1
                if i < len(lines): line = lines[i].strip()
            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for r_idx, row_cells in enumerate(table_data):
                    for c_idx, cell_text in enumerate(row_cells):
                        table.rows[r_idx].cells[c_idx].text = cell_text.replace('<br>', '\n')
            continue 

        else:
            line = re.sub(r'^\[\d+\]\s*', '', line)
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        i += 1
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- UI STREAMLIT ---
st.set_page_config(page_title="Pintar Syariah AI", page_icon="⚖️", layout="wide")
st.title("⚖️ Artificial Intelligence Mahkamah Syariah (AIMS)")
st.divider()

tab_kes, tab_pengurusan, tab_carian = st.tabs(["🏛️ Mod 1: Analisis Kes Syariah (AP)", "📝 Mod 2: Kertas Kerja", "🌐 Mod 3: Carian Pintar Online"])

# ==========================================
# MOD 1: ANALISIS KES (DRAF AP PRO)
# ==========================================
with tab_kes:
    col_meta, col_input = st.columns([1, 2]) 
    
    with col_meta:
        st.subheader("📋 Maklumat Kes (Metadata)")
        m_negeri = st.selectbox("Bidang Kuasa (Negeri):", ["Selangor", "Wilayah Persekutuan", "Johor", "Perak", "Kedah", "Kelantan", "Terengganu", "Pahang", "Negeri Sembilan", "Melaka", "Pulau Pinang", "Perlis", "Sabah", "Sarawak"])
        m_level = st.selectbox("Hierarki Mahkamah:", ["Mahkamah Rendah Syariah", "Mahkamah Tinggi Syariah", "Mahkamah Rayuan Syariah"])
        m_hakim = st.text_input("Nama Hakim / Panel:", placeholder="Cth: YA Tuan Haji...")
        m_tarikh = st.text_input("Tarikh Sidang / Keputusan:", placeholder="Cth: 26 Ramadhan 1438H / 21 Jun 2017")
        m_nokes = st.text_input("No. Kes:", placeholder="Cth: 10000-003-0001-2024")
        m_pihak1 = st.text_area("Pihak 1 (Plaintif/Pemohon):", height=70)
        m_pihak2 = st.text_area("Pihak 2 (Defendan/Responden):", height=70)
        m_jenis = st.text_area("Jenis Permohonan (Peruntukan):", placeholder="Cth: Permohonan Semakan di bawah seksyen 68...", height=70)
        m_peguam = st.text_input("Nama Peguam (Jika ada):", placeholder="Cth: Mohd. Faiz b. Adnan...")

    with col_input:
        st.subheader("📝 Ringkasan Fakta & Seksyen Utama")
        f_permohonan = st.text_area("Butiran PERMOHONAN (Jika ada):", placeholder="Biarkan kosong jika mahu AI drafkan...")
        f_fakta = st.text_area("Butiran FAKTA KES (Wajib isi):", height=150)
        f_ulasan = st.text_area("Butiran ULASAN MAHKAMAH (Jika ada):")
        f_keputusan = st.text_area("Butiran KEPUTUSAN (Jika ada):")

        if st.button("🔍 Jana Draf Alasan Penghakiman (AP)", type="primary"):
            if f_fakta.strip() == "":
                st.warning("⚠️ Sila masukkan Fakta Kes!")
            else:
                with st.spinner("AI sedang menganalisa maklumat yang diberikan dan menyediakan cadangan AP..."):
                    try:
                        embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
                        db = PineconeVectorStore(index_name="projek-aims", embedding=embeddings)
                        
                        dokumen_relevan = db.similarity_search(f_fakta, k=15)
                        konteks = "\n".join([d.page_content for d in dokumen_relevan])

                        prompt_ap = f"""Anda adalah seorang Hakim Mahkamah Syariah yang pakar. Tugas anda merangka ALASAN PENGHAKIMAN (AP) yang terperinci dan panjang.

AMARAN KERAS FORMAT:
1. JANGAN tulis Kepala Surat sama sekali. Mula draf terus dengan tajuk: PERMOHONAN
2. DILARANG menggunakan simbol Markdown (*, >, #, ##).
3. DILARANG menggunakan kurungan bernombor seperti [1], [2] pada mula perenggan.

PENTING UNTUK ULASAN MAHKAMAH: 
Anda DIWAJIBKAN memasukkan elemen ini dari Konteks:
- Peruntukan Undang-Undang: Petik nama akta/enakmen dan seksyen yang tepat bagi {m_negeri}.
- Autoriti Kes Lepas: Anda WAJIB memetik SEKURANG-KURANGNYA 3 HINGGA 4 KES LEPAS YANG BERBEZA daripada rujukan konteks di bawah untuk menyokong hujah. Terangkan kaitan setiap kes tersebut.
- Nas Syarak: Masukkan petikan hukum syarak berkaitan.

STRUKTUR KANDUNGAN:
PERMOHONAN
({f_permohonan})

FAKTA KES
({f_fakta})

ULASAN MAHKAMAH
(Ini adalah teras penghakiman. Huraikan kaitan fakta kes dengan undang-undang, 3 HINGGA 4 kes autoriti, dan hukum syarak. Ringkasan asas: {f_ulasan})

KEPUTUSAN
(Mesti dimulakan dengan: "SETELAH Kami membaca dan meneliti permohonan..." ATAU "SETELAH Mahkamah meneliti...")
(Guna laras bahasa hierarki {m_level}. Ringkasan: {f_keputusan})

RUJUKAN KES LEPAS (KONTEKS DARI PINECONE):
---------------------
{konteks}
---------------------
"""
                        respons = cuba_jana_ai(prompt_ap)
                        
                        meta_dict = {
                            'negeri': m_negeri, 'mahkamah': m_level, 'hakim': m_hakim, 'tarikh': m_tarikh, 
                            'nokes': m_nokes, 'pihak1': m_pihak1, 'pihak2': m_pihak2, 
                            'jenis_p': m_jenis, 'peguam': m_peguam
                        }
                        
                        fail_word = bina_fail_word(respons.text, "ALASAN PENGHAKIMAN", meta_dict)
                        
                        st.success("✅ Draf selesai dijana! Sila muat turun fail Word di bawah.")
                        st.download_button("📄 Muat Turun AP (Word)", data=fail_word, file_name=f"AP_{m_nokes}.docx")
                    except Exception as e: st.error(f"Ralat: {e}")

# ==========================================
# MOD 2: KERTAS KERJA PENGURUSAN
# ==========================================
with tab_pengurusan:
    st.info("Kertas Kerja Pengurusan & Kertas Konsep Arahan Amalan")
    jenis_kertas = st.selectbox("Jenis Dokumen:", ["Kertas Kerja Bengkel / Program", "Kertas Kerja Bajet", "Kertas Konsep Arahan Amalan"])
    
    if jenis_kertas == "Kertas Konsep Arahan Amalan":
        bahagian_unit = st.text_input("1. Bahagian / Unit Penyedia:", value="Bahagian Dasar & Penyelidikan (BPKR)")
        nama_program = st.text_input("2. Tajuk Kertas Konsep:", placeholder="Cth: Penangguhan Kes Atas Alasan Sijil Cuti Sakit")
        col1, col2 = st.columns(2)
        with col1: latar_belakang = st.text_area("3. Latar Belakang / Ringkasan Isu:", height=150)
        with col2: asas_pertimbangan = st.text_area("4. Asas Pertimbangan:", height=150)
        cadangan_syor = st.text_area("5. Cadangan & Syor:", height=100)
    else:
        bahagian_unit = st.text_input("1. Bahagian / Unit Penyedia:")
        nama_program = st.text_input("2. Nama Program / Aktiviti:")
        tarikh_masa = st.text_input("3. Tarikh, Masa & Tempoh:")
        tempat_program = st.text_input("4. Tempat / Lokasi Program:")
        col1, col2 = st.columns(2)
        with col1:
            maklumat_peserta = st.text_area("5. Maklumat & Bilangan Peserta:", height=150)
            maklumat_penceramah = st.text_area("6. Maklumat Penceramah:", height=150)
        with col2:
            kos_makan_minum = st.text_area("7. Anggaran Kos Kewangan:", height=150)
            objektif_tambahan = st.text_area("8. Objektif & Latar Belakang Tambahan:", height=150)

    if st.button("🚀 Jana Kertas Cadangan Lengkap", type="primary", key="btn_pengurusan"):
        if nama_program.strip() == "" or bahagian_unit.strip() == "":
            st.warning("⚠️ Sila isi Bahagian/Unit dan Tajuk terlebih dahulu bro!")
        else:
            with st.spinner(f"Menyusun format {jenis_kertas} rasmi melalui Pinecone..."):
                try:
                    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
                    db = PineconeVectorStore(index_name="projek-aims", embedding=embeddings)
                    filter_kategori = "Arahan Amalan" if jenis_kertas == "Kertas Konsep Arahan Amalan" else "Pengurusan"
                    dokumen_relevan = db.similarity_search(nama_program, k=5, filter={"sumber": filter_kategori})
                    konteks_teks = "\n".join([f"\n--- CONTOH TEMPLAT {idx+1} ---\n{doc.page_content}\n" for idx, doc in enumerate(dokumen_relevan)])

                    if jenis_kertas == "Kertas Konsep Arahan Amalan":
                        prompt_pengurusan = f"Bina draf Kertas Konsep rasmi. Tujuan, Latar Belakang: {latar_belakang}, Asas: {asas_pertimbangan}, Cadangan: {cadangan_syor}. Format dari: {konteks_teks}"
                    else:
                        prompt_pengurusan = f"Bina {jenis_kertas}. Unit: {bahagian_unit}, Latar: {objektif_tambahan}, Maklumat: {tarikh_masa}, Tempat: {tempat_program}. Kewangan: {kos_makan_minum}. Format dari: {konteks_teks}. WAJIB buat jadual dengan kolum Bil | Butiran | Jumlah."
                    
                    respons = cuba_jana_ai(prompt_pengurusan)
                    st.divider()
                    st.subheader(f"💡 Hasil Penjanaan AI ({jenis_kertas})")
                    st.write(respons.text)

                    meta_dummy = {'mahkamah':'','hakim':'','tarikh':'','nokes':'','pihak1':'','pihak2':'','jenis_p':'','peguam':'','negeri':''}
                    is_aa = (jenis_kertas == "Kertas Konsep Arahan Amalan")
                    tajuk_rasmi = f"KERTAS KONSEP ARAHAN AMALAN\n{nama_program.upper()}" if is_aa else f"KERTAS PERMOHONAN KELULUSAN BERBELANJA BAGI\n{nama_program.upper()}\nJABATAN KEHAKIMAN SYARIAH MALAYSIA"
                    
                    fail_pengurusan_docx = bina_fail_word(respons.text, tajuk_rasmi, meta_dummy, is_pengurusan=True, is_arahan_amalan=is_aa)
                    st.download_button(f"📄 Muat Turun {jenis_kertas} (Word)", data=fail_pengurusan_docx, file_name=f"{jenis_kertas.replace(' ', '_')}.docx")
                except Exception as e: st.error(f"❌ Ralat Sistem: {e}")

# ==========================================
# MOD 3: CARIAN PINTAR ONLINE (DENGAN BUTANG CARIAN LANJUT)
# ==========================================
with tab_carian:
    st.subheader("🌐 Carian Maklumat Pintar (Online)")
    st.write("Sistem ini akan mencari maklumat perundangan atau maklumat umum terkini di internet (melalui DuckDuckGo) dan AI akan merumuskannya untuk anda.")
    
    # 1. SETUP SESSION STATE
    if "hasil_asas" not in st.session_state:
        st.session_state.hasil_asas = ""
    if "carian_query" not in st.session_state:
        st.session_state.carian_query = ""

    topik_carian = st.text_input("🔍 Masukkan isu atau topik yang ingin dicari (Cth: Apa itu permohonan cerai fasakh mengikut undang-undang keluarga Islam?):")
    
    # 2. BUTANG PERTAMA (Carian Asas)
    if st.button("Jalankan Carian", type="primary"):
        if topik_carian.strip() == "":
            st.warning("⚠️ Sila masukkan topik carian terlebih dahulu.")
        else:
            with st.spinner("Sedang mencari di internet dan merumus maklumat..."):
                try:
                    carian_enjin = DuckDuckGoSearchRun()
                    hasil_mentah = carian_enjin.invoke(topik_carian)
                    
                    prompt_carian = f"""Anda adalah Pembantu Penyelidik Syariah. 
                    Berikut adalah maklumat mentah yang diperolehi dari internet mengenai topik: '{topik_carian}'.
                    Tugas anda adalah menyusun semula maklumat ini menjadi satu penerangan yang sangat kemas, profesional, dan mudah difahami.
                    
                    Maklumat Mentah:
                    {hasil_mentah}
                    """
                    respons_carian = cuba_jana_ai(prompt_carian)
                    
                    # Simpan dalam session state supaya tak hilang
                    st.session_state.hasil_asas = respons_carian.text
                    st.session_state.carian_query = topik_carian
                    
                except Exception as e:
                    st.error(f"❌ Gagal melakukan carian: {e}")

    # 3. PAPARAN HASIL PERTAMA & BUTANG KEDUA
    if st.session_state.hasil_asas:
        st.success("✅ Carian Berjaya! Anda boleh rujuk ringkasan asas di bawah:")
        st.text_area("Hasil Carian", value=st.session_state.hasil_asas, height=350)
        
        st.divider()
        
        # Butang Carian Lanjut
        if st.button("🔍 Carian Lanjut (Akta, Kes & Link)"):
            with st.spinner("Sedang menggali peruntukan undang-undang dan kes rujukan beserta pautan..."):
                try:
                    carian_enjin = DuckDuckGoSearchRun()
                    query_lanjut = f"{st.session_state.carian_query} akta enakmen kes mahkamah syariah e-syariah malaysia pautan link"
                    hasil_mentah_lanjut = carian_enjin.invoke(query_lanjut)
                    
                    prompt_lanjut = f"""
                    Anda adalah Penyelidik Undang-Undang Syariah Kanan. Berdasarkan isu: '{st.session_state.carian_query}' 
                    dan hasil carian web ini: '{hasil_mentah_lanjut}'.
                    
                    Sila sediakan maklumat terperinci berikut:
                    1. **Peruntukan Undang-Undang**: Senaraikan Akta/Enakmen Syariah yang berkaitan di Malaysia.
                    2. **Kes Rujukan**: Senarai kes-kes Mahkamah Syariah yang telah diputuskan berhubung isu ini.
                    3. **Pautan (URL)**: Anda WAJIB menyertakan pautan (URL link) internet yang sah yang dijumpai dalam carian web tersebut supaya pengguna boleh klik untuk bacaan lanjut.
                    
                    Susun jawapan dengan kemas menggunakan format Markdown (Contoh format: `[Nama Kes/Akta](URL)`). 
                    Jika tiada pautan khusus dijumpai dalam hasil carian, berikan cadangan portal rasmi (seperti portal e-Syariah atau Jurnal Hukum).
                    """
                    
                    respons_lanjut = cuba_jana_ai(prompt_lanjut)
                    
                    st.markdown("### 📚 Hasil Carian Lanjut (Perundangan, Kes & Pautan)")
                    st.info("Nota: Pautan (URL) yang dijana bergantung kepada ketersediaan rekod di enjin carian awam internet.")
                    st.markdown(respons_lanjut.text)
                    
                except Exception as e:
                    st.error(f"❌ Gagal melakukan carian lanjut: {e}")

# ==========================================
# FOOTER HAK CIPTA (HAK MILIK EKSKLUSIF)
# ==========================================
footer_html = """
<style>
.footer {
    position: fixed;
    left: 0;
    bottom: 0;
    width: 100%;
    background-color: #0F172A;
    color: #94A3B8;
    text-align: center;
    padding: 12px;
    font-size: 13px;
    border-top: 4px solid #1D4ED8;
    z-index: 999;
}
.block-container { padding-bottom: 80px; }
</style>
<div class="footer">
    © 2026 AIMS Dibangunkan Oleh Pihak UTES. Hak Cipta Terpelihara.
</div>
"""
st.markdown(footer_html, unsafe_allow_html=True)