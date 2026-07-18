import os
import time
import pypdf
from dotenv import load_dotenv
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain_core.documents import Document as LangchainDocument

# 1. Muatkan Kunci API
load_dotenv()
index_name = os.getenv("PINECONE_INDEX_NAME")

# 2. Inisialisasi OpenAI Embeddings
embeddings = OpenAIEmbeddings(model="text-embedding-3-small")

# --- FUNGSI BACA WORD ---
def baca_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])

# --- FUNGSI BACA PDF ---
def baca_pdf(file_path):
    teks_penuh = ""
    with open(file_path, "rb") as fail:
        pembaca = pypdf.PdfReader(fail)
        for muka_surat in pembaca.pages:
            teks = muka_surat.extract_text()
            if teks:
                teks_penuh += teks + "\n"
    return teks_penuh

# --- FUNGSI REKOD FAIL YANG TELAH DISERAP ---
FAIL_REKOD = "rekod_serapan.txt"

def dapatkan_rekod_lama():
    """Baca senarai nama fail yang dah pernah dimasukkan ke Pinecone."""
    if os.path.exists(FAIL_REKOD):
        with open(FAIL_REKOD, "r", encoding="utf-8") as f:
            return set(f.read().splitlines())
    return set()

def simpan_rekod_baru(senarai_fail):
    """Simpan nama fail-fail baru yang dah berjaya masuk Pinecone ke dalam rekod."""
    with open(FAIL_REKOD, "a", encoding="utf-8") as f:
        for fail in senarai_fail:
            f.write(fail + "\n")

# 3. Proses Membaca Dokumen (Selongkar Subfolder & Elak Error & Elak Duplikasi)
folder_data = "data_kes"
semua_teks = []
fail_telah_diserap = dapatkan_rekod_lama()
fail_baru_untuk_direkod = []

print(f"Mencari fail dalam folder: '{folder_data}' dan semua subfoldernya...")
if not os.path.exists(folder_data):
    os.makedirs(folder_data)
    print(f"Folder '{folder_data}' baru dicipta. Sila susun fail dan folder ke dalamnya.")
else:
    for root, dirs, files in os.walk(folder_data):
        for filename in files:
            
            # --- SEMAKAN REKOD LAMA (ELAK PERTINDIHAN) ---
            if filename in fail_telah_diserap:
                print(f"⏭️  Abaikan (Dah wujud dalam Pinecone): {filename}")
                continue

            laluan_fail = os.path.join(root, filename)
            teks = ""
            
            # --- BLOK TRY-EXCEPT UNTUK ELAK CRASH ---
            try:
                # Semak jenis fail dan baca
                if filename.lower().endswith(".docx"):
                    print(f"📄 Membaca Word: {os.path.relpath(laluan_fail, folder_data)}")
                    teks = baca_docx(laluan_fail)
                elif filename.lower().endswith(".pdf"):
                    print(f"📑 Membaca PDF: {os.path.relpath(laluan_fail, folder_data)}")
                    teks = baca_pdf(laluan_fail)
                    
                # Jika ada teks, simpan bersama metadatanya
                if teks.strip():
                    semua_teks.append(LangchainDocument(
                        page_content=teks, 
                        metadata={"sumber": filename, "kategori": os.path.basename(root)}
                    ))
                    # Masukkan ke dalam senarai berasingan untuk direkodkan selepas berjaya
                    fail_baru_untuk_direkod.append(filename)
            
            except Exception as e:
                print(f"⚠️ GAGAL MEMBACA FAIL: {filename} | Ralat: {e}")
                continue 

    # 4. Semak sama ada ada fail baru atau tidak
    if not semua_teks:
        print("\n👍 Tiada dokumen baru untuk diserap. Semua sedia ada dah up-to-date!")
    else:
        # Potong Teks (Chunking)
        print("\nMemproses teks (Chunking)...")
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        docs = text_splitter.split_documents(semua_teks)
        print(f"Jumlah fail baru dibaca: {len(fail_baru_untuk_direkod)} | Jumlah pecahan (chunks): {len(docs)}")

        # 5. Serap ke Pinecone (Kaedah Kura-kura)
        batch_size = 20
        vectorstore = PineconeVectorStore(index_name=index_name, embedding=embeddings)

        print("\n🚀 Mula memuat naik ke Pinecone secara berperingkat (Kaedah Kura-kura)...")
        for i in range(0, len(docs), batch_size):
            batch = docs[i:i+batch_size]
            vectorstore.add_documents(batch)
            
            pusingan = (i // batch_size) + 1
            jumlah_pusingan = (len(docs) // batch_size) + (1 if len(docs) % batch_size != 0 else 0)
            print(f"Berjaya muat naik kumpulan {pusingan}/{jumlah_pusingan}. Berehat 10 saat...")
            time.sleep(10) # Rehat 10 saat untuk elak 'Rate Limit'

        # 6. Simpan Rekod Fail Baru Ke Dalam rekod_serapan.txt
        simpan_rekod_baru(fail_baru_untuk_direkod)
        print("\nAlhamdulillah! Selesai serap dokumen baru ke Pinecone dan rekod telah dikemaskini.")